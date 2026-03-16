"""Export proposal to Markdown, Word (.docx), and JSON formats.

Also provides helpers to create timestamped output directories and
save intermediate pipeline artifacts.
"""

from __future__ import annotations

import json
import logging
import os
import re
from datetime import datetime
from typing import Optional

from .models import ProposalPackage

logger = logging.getLogger(__name__)


def create_output_dir(
    rfp_name: str = "proposal",
    base_dir: str = "proposals",
) -> str:
    """Create a timestamped output directory for a proposal run.

    Structure: proposals/<rfp_name>_YYYY-MM-DD_HHMMSS/

    Args:
        rfp_name: Short name derived from the RFP (sanitized for filesystem).
        base_dir: Parent directory for all proposals.

    Returns:
        Absolute path to the created directory.
    """
    # Sanitize rfp_name for filesystem
    safe_name = re.sub(r"[^\w\s-]", "", rfp_name)[:60].strip().replace(" ", "_")
    if not safe_name:
        safe_name = "proposal"

    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    dir_name = f"{safe_name}_{timestamp}"
    full_path = os.path.join(base_dir, dir_name)
    os.makedirs(full_path, exist_ok=True)
    logger.info("Created output directory: %s", full_path)
    return full_path


def save_intermediate(data, filename: str, output_dir: str) -> str:
    """Save an intermediate pipeline artifact as JSON.

    Args:
        data: A Pydantic model or dict to serialize.
        filename: e.g. 'step1_rfp_analysis.json'
        output_dir: Directory to save into.

    Returns:
        Path to the saved file.
    """
    path = os.path.join(output_dir, filename)
    if hasattr(data, "model_dump_json"):
        content = data.model_dump_json(indent=2)
    elif hasattr(data, "model_dump"):
        content = json.dumps(data.model_dump(), indent=2, default=str)
    elif isinstance(data, dict):
        content = json.dumps(data, indent=2, default=str)
    elif isinstance(data, list):
        # Handle list of Pydantic models or dicts
        items = []
        for item in data:
            if hasattr(item, "model_dump"):
                items.append(item.model_dump())
            else:
                items.append(item)
        content = json.dumps(items, indent=2, default=str)
    else:
        content = json.dumps(str(data), indent=2)

    with open(path, "w") as f:
        f.write(content)
    logger.info("Saved intermediate: %s", path)
    return path


def to_json(package: ProposalPackage) -> str:
    """Export the full proposal package as a JSON string."""
    return package.model_dump_json(indent=2)


def to_markdown(package: ProposalPackage) -> str:
    """Export the full proposal as a single markdown string.

    Args:
        package: The completed proposal package.

    Returns:
        Full proposal content as markdown text.
    """
    parts = []

    # Title
    parts.append(f"# Proposal for {package.rfp_analysis.customer}")
    parts.append("")

    # Metadata
    meta = package.generation_metadata
    parts.append(
        f"*Generated using {meta.get('model', 'unknown')} | "
        f"{meta.get('total_words', 0):,} words | "
        f"Score: {package.overall_score:.0f}/100*"
    )
    parts.append("")
    parts.append("---")
    parts.append("")

    # Table of Contents
    parts.append("## Table of Contents")
    parts.append("")
    for section in package.sections:
        indent = "  " * (section.number.count("."))
        parts.append(f"{indent}- {section.number}. {section.title}")
    parts.append("")
    parts.append("---")
    parts.append("")

    # Sections
    for section in package.sections:
        level = section.number.count(".") + 2  # ## for top-level, ### for sub
        heading = "#" * min(level, 4)
        parts.append(f"{heading} {section.number}. {section.title}")
        parts.append("")
        parts.append(section.content)
        parts.append("")
        parts.append("---")
        parts.append("")

    # Compliance Matrix
    parts.append("## Compliance Matrix")
    parts.append("")
    parts.append(
        f"*Coverage: {package.compliance_matrix.coverage_percentage:.0f}% | "
        f"Gaps: {package.compliance_matrix.gap_count}*"
    )
    parts.append("")
    parts.append("| ID | Requirement | Section | Status |")
    parts.append("|---|---|---|---|")
    for row in package.compliance_matrix.rows:
        desc = row.requirement_description[:80]
        if len(row.requirement_description) > 80:
            desc += "..."
        parts.append(
            f"| {row.requirement_id} | {desc} | {row.proposal_section} | {row.compliance_status} |"
        )
    parts.append("")

    return "\n".join(parts)


def to_docx(package: ProposalPackage, output_path: str) -> str:
    """Export the proposal as a formatted Word document.

    Args:
        package: The completed proposal package.
        output_path: File path for the output .docx file.

    Returns:
        The output file path.
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Technical and Management Proposal")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Prepared for {package.rfp_analysis.customer}")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    meta_text = (
        f"Total Words: {package.generation_metadata.get('total_words', 0):,}\n"
        f"Overall Score: {package.overall_score:.0f}/100\n"
        f"Sections: {len(package.sections)}"
    )
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_para.add_run(meta_text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── Table of Contents placeholder ──
    toc_heading = doc.add_heading("Table of Contents", level=1)
    for s in package.sections:
        indent = "    " * s.number.count(".")
        doc.add_paragraph(
            f"{indent}{s.number}. {s.title}",
            style="List Number",
        )
    doc.add_page_break()

    # ── Proposal Sections ──
    for s in package.sections:
        level = min(s.number.count(".") + 1, 3)  # Heading 1-3
        doc.add_heading(f"{s.number}. {s.title}", level=level)

        # Parse markdown content into paragraphs
        _add_markdown_content(doc, s.content)

        # Add word count footer for each section
        wc_para = doc.add_paragraph()
        run = wc_para.add_run(f"[{s.word_count} words]")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.italic = True

    doc.add_page_break()

    # ── Compliance Matrix ──
    doc.add_heading("Compliance Matrix", level=1)
    doc.add_paragraph(
        f"Coverage: {package.compliance_matrix.coverage_percentage:.0f}% | "
        f"Gaps: {package.compliance_matrix.gap_count}"
    )

    if package.compliance_matrix.rows:
        table = doc.add_table(
            rows=1 + len(package.compliance_matrix.rows),
            cols=4,
        )
        table.style = "Light Grid Accent 1"

        # Header row
        headers = ["ID", "Requirement", "Section", "Status"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for i, row in enumerate(package.compliance_matrix.rows, 1):
            table.cell(i, 0).text = row.requirement_id
            desc = row.requirement_description
            if len(desc) > 100:
                desc = desc[:100] + "..."
            table.cell(i, 1).text = desc
            table.cell(i, 2).text = row.proposal_section
            table.cell(i, 3).text = row.compliance_status

    doc.add_page_break()

    # ── Scoring Summary ──
    doc.add_heading("Proposal Scoring Summary", level=1)
    doc.add_paragraph(f"Overall Score: {package.overall_score:.0f}/100")

    if package.scores:
        score_table = doc.add_table(
            rows=1 + len(package.scores),
            cols=4,
        )
        score_table.style = "Light Grid Accent 1"

        score_headers = ["Section", "Score", "Strengths", "Weaknesses"]
        for i, header in enumerate(score_headers):
            cell = score_table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for i, score in enumerate(package.scores, 1):
            score_table.cell(i, 0).text = (
                f"{score.section_number}. {score.section_title}"
            )
            score_table.cell(i, 1).text = str(score.score)
            score_table.cell(i, 2).text = "; ".join(score.strengths[:2])
            score_table.cell(i, 3).text = "; ".join(score.weaknesses[:2])

    # ── Footer with page numbers ──
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Save
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    logger.info("Saved .docx to %s", output_path)
    return output_path


def _add_markdown_content(doc, content: str):
    """Parse basic markdown content and add to the Word document."""
    from docx.shared import Pt

    lines = content.split("\n")
    current_paragraph_lines = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            # Flush current paragraph
            if current_paragraph_lines:
                doc.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            if current_paragraph_lines:
                doc.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        # Numbered lists
        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in (".", ")"):
            if current_paragraph_lines:
                doc.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            doc.add_paragraph(stripped[2:].strip(), style="List Number")
            continue

        # Sub-headings within section content
        if stripped.startswith("###"):
            if current_paragraph_lines:
                doc.add_paragraph(" ".join(current_paragraph_lines))
                current_paragraph_lines = []
            doc.add_heading(stripped.lstrip("#").strip(), level=3)
            continue

        # Regular text — accumulate into paragraph
        current_paragraph_lines.append(stripped)

    # Flush remaining
    if current_paragraph_lines:
        doc.add_paragraph(" ".join(current_paragraph_lines))
